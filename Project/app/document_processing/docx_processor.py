import docx
import os
import logging
from typing import Dict, List, Any, Optional
import pytesseract
from PIL import Image
import io
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter

from .processor import BaseDocumentProcessor, Document, DocumentChunk
from ..core.config import settings

logger = logging.getLogger(__name__)

class DocxProcessor(BaseDocumentProcessor):
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len
        )
    
    def process(self, file_path: str) -> Document:
        """
        Process a DOCX file and extract text and metadata.
        For images embedded in the document, OCR will be applied.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Document: A Document object with extracted text and metadata
        """
        logger.info(f"Processing DOCX file: {file_path}")
        
        try:
            # Load the document
            doc = docx.Document(file_path)
            
            # Extract metadata
            doc_metadata = {
                "source": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": "docx",
                "page_count": self._estimate_page_count(doc),
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "creation_date": str(doc.core_properties.created) if doc.core_properties.created else "",
            }
            
            # Extract text from paragraphs
            text_content = []
            
            # Estimate page breaks and add page markers
            pages = self._estimate_pages(doc)
            
            # Process document content with page markers
            full_text = ""
            for i, (page_num, page_content) in enumerate(pages.items(), 1):
                full_text += f"--- Page {page_num} ---\n{page_content}\n\n"
            
            # Process images in the document for OCR
            image_text = self._process_images(doc)
            if image_text:
                full_text += f"\n--- Images OCR Text ---\n{image_text}"
            
            return Document(full_text, doc_metadata)
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise
    
    def _estimate_page_count(self, doc: docx.Document) -> int:
        """Estimate the number of pages in a DOCX document."""
        # A very rough estimation: assume 500 words per page
        words_per_page = 500
        total_words = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
        return max(1, round(total_words / words_per_page))
    
    def _estimate_pages(self, doc: docx.Document) -> Dict[int, str]:
        """
        Estimate page breaks in the document and return text by page.
        This is an approximation since docx doesn't have direct page information.
        """
        # Approximate characters per page (based on A4 page with normal margins and font)
        chars_per_page = 3000
        pages = {}
        current_page = 1
        current_page_text = ""
        
        for paragraph in doc.paragraphs:
            current_page_text += paragraph.text + "\n"
            
            # Check if we've exceeded the estimated page length
            if len(current_page_text) > chars_per_page:
                pages[current_page] = current_page_text
                current_page += 1
                current_page_text = ""
        
        # Add the last page if there's content
        if current_page_text:
            pages[current_page] = current_page_text
        
        # If no pages were created, add at least one
        if not pages:
            pages[1] = ""
            
        return pages
    
    def _process_images(self, doc: docx.Document) -> str:
        """Extract text from images in the document using OCR."""
        image_texts = []
        
        # Try to extract images - this is simplified and might not work for all DOCXs
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_blob = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_blob))
                        text = pytesseract.image_to_string(image, lang=settings.OCR_LANGUAGE)
                        if text and not text.isspace():
                            image_texts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to process image: {e}")
        except Exception as e:
            logger.warning(f"Failed to extract images from document: {e}")
        
        return "\n\n".join(image_texts)
    
    def chunk(self, document: Document) -> List[DocumentChunk]:
        """
        Split a document into chunks for embedding.
        
        Args:
            document: Document object to chunk
            
        Returns:
            List[DocumentChunk]: List of document chunks
        """
        logger.info(f"Chunking document: {document.metadata.get('source', 'unknown')}")
        
        try:
            # Split the document text into chunks
            chunks = self.text_splitter.split_text(document.text)
            
            # Create DocumentChunk objects with metadata
            doc_chunks = []
            for i, chunk_text in enumerate(chunks):
                # Extract page numbers from chunk text
                page_numbers = []
                page_pattern = re.compile(r"--- Page (\d+) ---")
                for match in page_pattern.finditer(chunk_text):
                    page_numbers.append(int(match.group(1)))
                
                # If no page markers found, make an estimate
                if not page_numbers:
                    # Simplified estimation
                    total_length = len(document.text)
                    chunk_start = document.text.find(chunk_text)
                    relative_position = chunk_start / total_length if total_length > 0 else 0
                    estimated_page = max(1, min(
                        round(relative_position * document.metadata.get("page_count", 1)), 
                        document.metadata.get("page_count", 1)
                    ))
                    page_numbers = [estimated_page]
                
                # Create chunk metadata
                chunk_metadata = document.metadata.copy()
                chunk_metadata.update({
                    "chunk_id": i,
                    "pages": page_numbers,
                    "page": page_numbers[0] if page_numbers else 1,  # For compatibility with the API
                })
                
                doc_chunks.append(DocumentChunk(chunk_text, chunk_metadata))
            
            return doc_chunks
            
        except Exception as e:
            logger.error(f"Error chunking document: {e}")
            raise